import json
from typing import List, Dict, Any, Optional
from openpyxl import load_workbook
from docx import Document


class QuestionImportError(Exception):
    def __init__(self, row: int, message: str):
        self.row = row
        self.message = message
        super().__init__(f"第{row}行: {message}")


def parse_options(options_str: str) -> List[Dict[str, str]]:
    if not options_str:
        return []
    
    options = []
    parts = options_str.split('|')
    for part in parts:
        part = part.strip()
        if not part:
            continue
        if ':' in part:
            opt_id, content = part.split(':', 1)
            options.append({"id": opt_id.strip(), "content": content.strip()})
        else:
            opt_id = chr(65 + len(options))
            options.append({"id": opt_id, "content": part})
    return options


def parse_judgment_options(options_str: str) -> List[Dict[str, str]]:
    if not options_str:
        return [{"id": "true", "content": "正确"}, {"id": "false", "content": "错误"}]
    
    parts = [p.strip() for p in options_str.split('|') if p.strip()]
    if len(parts) >= 2:
        return [{"id": "true", "content": parts[0]}, {"id": "false", "content": parts[1]}]
    return [{"id": "true", "content": "正确"}, {"id": "false", "content": "错误"}]


def validate_question_type(qtype: str) -> str:
    qtype = qtype.strip().lower()
    if qtype in ['single', '单选', '单选题']:
        return 'single'
    elif qtype in ['multiple', '多选', '多选题']:
        return 'multiple'
    elif qtype in ['judgment', '判断', '判断题']:
        return 'judgment'
    elif qtype in ['shared_option', '共用选项', '共用选项题']:
        return 'shared_option'
    raise ValueError(f"未知题型: {qtype}")


def parse_single_row(
    row_data: List[str],
    row_num: int
) -> Dict[str, Any]:
    if len(row_data) < 4:
        raise QuestionImportError(row_num, "数据列数不足，至少需要：题型、题干、选项、答案")

    question_type = validate_question_type(row_data[0])
    content = row_data[1].strip() if len(row_data) > 1 else ""
    options_str = row_data[2].strip() if len(row_data) > 2 else ""
    answer = row_data[3].strip() if len(row_data) > 3 else ""
    explanation = row_data[4].strip() if len(row_data) > 4 else ""
    score = int(row_data[5].strip()) if len(row_data) > 5 and row_data[5].strip() else 10

    if not content:
        raise QuestionImportError(row_num, "题干不能为空")
    if not answer:
        raise QuestionImportError(row_num, "答案不能为空")

    if question_type == 'judgment':
        options = parse_judgment_options(options_str)
        answer = "true" if answer in ["正确", "true", "True", "1", "T"] else "false"
    elif question_type == 'shared_option':
        # 共用选项题不需要选项，只存储空数组
        options = []
        if ',' in answer:
            answer = answer.split(',')[0]
    else:
        options = parse_options(options_str)
        if question_type == 'single' and ',' in answer:
            answer = answer.split(',')[0]
        if question_type == 'multiple':
            answer = ','.join([a.strip() for a in answer.replace(',', '|').split('|') if a.strip()])

    return {
        "question_type": question_type,
        "content": content,
        "options": options,
        "answer": answer,
        "explanation": explanation,
        "score": score
    }


def import_from_excel(file_content: bytes) -> List[Dict[str, Any]]:
    import io
    wb = load_workbook(io.BytesIO(file_content), data_only=True)
    ws = wb.active
    
    questions = []
    errors = []
    
    header_row = True
    for row_num, row in enumerate(ws.iter_rows(values_only=True), start=1):
        if header_row:
            header_row = False
            continue
        
        if not row or not any(row):
            continue
        
        try:
            row_data = [str(cell) if cell is not None else "" for cell in row]
            question = parse_single_row(row_data, row_num)
            questions.append(question)
        except QuestionImportError as e:
            errors.append({"row": e.row, "error": e.message})
        except Exception as e:
            errors.append({"row": row_num, "error": str(e)})
    
    return questions, errors


def import_from_docx(file_content: bytes) -> List[Dict[str, Any]]:
    import io
    doc = Document(io.BytesIO(file_content))
    
    questions = []
    errors = []
    
    table = None
    for element in doc.element.body:
        if element.tag.endswith('}tbl'):
            table = doc.tables[len([t for t in doc.tables if t._element < element]) - 1] if doc.tables else None
            break
    
    if not table:
        for t in doc.tables:
            table = t
            break
    
    if not table:
        raise ValueError("Word文档中未找到表格")
    
    header_row = True
    for row_idx, row in enumerate(table.rows, start=1):
        if header_row:
            header_row = False
            continue
        
        row_data = [cell.text.strip() for cell in row.cells]
        
        if not any(row_data):
            continue
        
        try:
            question = parse_single_row(row_data, row_idx)
            questions.append(question)
        except QuestionImportError as e:
            errors.append({"row": e.row, "error": e.message})
        except Exception as e:
            errors.append({"row": row_idx, "error": str(e)})
    
    return questions, errors


def import_from_txt(file_content: bytes, delimiter: str = ",") -> List[Dict[str, Any]]:
    import io
    # Try multiple encodings to support Chinese files
    encodings = ['utf-8', 'gbk', 'gb2312', 'gb18030']
    content = None
    for encoding in encodings:
        try:
            content = file_content.decode(encoding)
            break
        except UnicodeDecodeError:
            continue

    if content is None:
        raise ValueError("无法解析文件编码，请使用 UTF-8、GBK、GB2312 或 GB18030 编码的文件")
    lines = content.strip().split('\n')
    
    questions = []
    errors = []
    
    for row_num, line in enumerate(lines, start=1):
        line = line.strip()
        if not line:
            continue
        
        row_data = [field.strip() for field in line.split(delimiter)]
        
        try:
            question = parse_single_row(row_data, row_num)
            questions.append(question)
        except QuestionImportError as e:
            errors.append({"row": e.row, "error": e.message})
        except Exception as e:
            errors.append({"row": row_num, "error": str(e)})
    
    return questions, errors


def import_questions(file_content: bytes, filename: str) -> List[Dict[str, Any]]:
    ext = filename.lower().split('.')[-1]

    if ext in ['xlsx', 'xls']:
        return import_from_excel(file_content)
    elif ext == 'docx':
        return import_from_docx(file_content)
    elif ext in ['txt', 'csv']:
        return import_from_txt(file_content)
    else:
        raise ValueError(f"不支持的文件格式: {ext}")


def import_shared_options(file_content: bytes, filename: str) -> List[Dict[str, Any]]:
    """Import shared option groups from Excel or CSV"""
    ext = filename.lower().split('.')[-1]

    if ext in ['xlsx', 'xls']:
        return import_shared_options_from_excel(file_content)
    elif ext in ['csv']:
        return import_shared_options_from_csv(file_content)
    else:
        raise ValueError(f"不支持的文件格式: {ext}. 请使用 .xlsx, .xls 或 .csv")


def import_shared_options_from_excel(file_content: bytes) -> List[Dict[str, Any]]:
    """Import shared option groups from Excel file"""
    import io
    wb = load_workbook(io.BytesIO(file_content), data_only=True)
    ws = wb.active

    groups = []
    errors = []

    header_row = True
    for row_num, row in enumerate(ws.iter_rows(values_only=True), start=1):
        if header_row:
            header_row = False
            continue

        if not row or not any(row):
            continue

        try:
            row_data = [str(cell) if cell is not None else "" for cell in row]

            if len(row_data) < 2:
                errors.append({"row": row_num, "error": "数据列数不足，至少需要：选项组名称、选项列表"})
                continue

            group_name = row_data[0].strip()
            options_str = row_data[1].strip() if len(row_data) > 1 else ""

            if not group_name:
                errors.append({"row": row_num, "error": "选项组名称不能为空"})
                continue

            # Parse options
            options = []
            if options_str:
                parts = [p.strip() for p in options_str.split('|') if p.strip()]
                for part in parts:
                    if ':' in part:
                        opt_id, content = part.split(':', 1)
                        options.append({"id": opt_id.strip(), "content": content.strip()})
                    else:
                        opt_id = chr(65 + len(options))
                        options.append({"id": opt_id, "content": part.strip()})

            if len(options) < 2:
                errors.append({"row": row_num, "error": "选项列表至少需要2个选项"})
                continue

            groups.append({
                "name": group_name,
                "options": options
            })
        except Exception as e:
            errors.append({"row": row_num, "error": str(e)})

    return groups, errors


def import_shared_options_from_csv(file_content: bytes) -> List[Dict[str, Any]]:
    """Import shared option groups from CSV file"""
    import io
    # Try multiple encodings to support Chinese files
    encodings = ['utf-8', 'gbk', 'gb2312', 'gb18030']
    content = None
    for encoding in encodings:
        try:
            content = file_content.decode(encoding)
            break
        except UnicodeDecodeError:
            continue

    if content is None:
        raise ValueError("无法解析文件编码，请使用 UTF-8、GBK、GB2312 或 GB18030 编码的文件")
    lines = content.strip().split('\n')

    groups = []
    errors = []

    header_row = True
    for row_num, line in enumerate(lines, start=1):
        if header_row:
            header_row = False
            continue

        line = line.strip()
        if not line:
            continue

        try:
            row_data = [field.strip() for field in line.split(',')]

            if len(row_data) < 2:
                errors.append({"row": row_num, "error": "数据列数不足，至少需要：选项组名称、选项列表"})
                continue

            group_name = row_data[0].strip()
            options_str = row_data[1].strip() if len(row_data) > 1 else ""

            if not group_name:
                errors.append({"row": row_num, "error": "选项组名称不能为空"})
                continue

            # Parse options
            options = []
            if options_str:
                parts = [p.strip() for p in options_str.split('|') if p.strip()]
                for part in parts:
                    if ':' in part:
                        opt_id, content = part.split(':', 1)
                        options.append({"id": opt_id.strip(), "content": content.strip()})
                    else:
                        opt_id = chr(65 + len(options))
                        options.append({"id": opt_id, "content": part.strip()})

            if len(options) < 2:
                errors.append({"row": row_num, "error": "选项列表至少需要2个选项"})
                continue

            groups.append({
                "name": group_name,
                "options": options
            })
        except Exception as e:
            errors.append({"row": row_num, "error": str(e)})

    return groups, errors
